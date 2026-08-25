"""Generate the EHS Commentary manuscript (docx).

Commentary format: ~1000 words, one figure, APA author-date references.
All numeric values are read from results/*.json (produced by s02-s06).
"""
import json
import os
import re

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from common import FIGS_DIR, RESULTS_DIR, ROOT

MANU_DIR = os.path.join(ROOT, "manuscript")

def _pct(x):
    return f"{100 * x:.0f}%"

def _load():
    j = lambda n: json.load(open(os.path.join(RESULTS_DIR, n)))
    base = j("baseline_calendar.json")
    onset = j("onset_summary.json")
    cp = j("changepoint_uncertainty.json")
    ev = j("event_alignment.json")
    sens = j("sensitivity_summary.json")
    sub = j("subnational_summary.json")
    hmd_c = j("subnational_hmd_contrast.json")
    subpts = pd.read_csv(os.path.join(RESULTS_DIR, "subnational_points.csv"))
    return base, onset, cp, ev, sens, sub, hmd_c, subpts

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_caption(doc, text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = True
    return p

def add_figure(doc, path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.add_run().add_picture(path, width=Inches(width))
    add_caption(doc, caption)

def _set_continuous_line_numbers(section):
    sectPr = section._sectPr
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:restart"), "continuous")
    lnNumType.set(qn("w:distance"), "720")
    sectPr.append(lnNumType)

def _add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.clear()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE \\* MERGEFORMAT "
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)
    run2 = p.add_run("1")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar3)

def _add_line_numbers_and_page_numbers(doc):
    for section in doc.sections:
        _set_continuous_line_numbers(section)
        _add_page_number_footer(section)

def _word_count(text):
    """Rough word count for English text, ignoring punctuation."""
    return len(re.findall(r"\b[\w'\u2019]+\b", text))

def main():
    os.makedirs(MANU_DIR, exist_ok=True)
    base, onset, cp, ev, sens, sub, hmd_c, subpts = _load()
    de = hmd_c["Germany"]
    boot = cp["global_change_year_bootstrap"]
    pco = cp["per_country_onset"]
    coll = ev["collapse"]
    cvc, cve = ev["conserved_cv_calendar"], ev["conserved_cv_event"]
    r2 = base["r2_by_k"]
    n_regions_total = int(subpts[subpts.source.isin(
        ["Eurostat NUTS2", "CDC/NCHS + US Census PEP"])].shape[0])
    matched = {k: v["lambda_cv_matched_e0"] for k, v in sub.items()
               if v.get("lambda_cv_matched_e0") is not None}
    matched_series = pd.Series(matched)
    es = sub["Spain"]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)

    # ---- Title ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = t.add_run("Two pathways, many clocks: a commentary on Itao (2026)")
    trun.bold = True
    trun.font.size = Pt(15)
    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    au.add_run("Tatsuki Onishi").font.size = Pt(12)
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = aff.add_run("Affiliation to be completed. Corresponding author: Tatsuki Onishi "
                     "(email to be completed).")
    ar.font.size = Pt(9)
    ar.italic = True

    # ---- Abstract ----
    H(doc, "Abstract", 2)
    abstract = (
        f"Itao (2026) reports two universal demographic-transition pathways: a conserved "
        f"{chr(955)}e{chr(8320)} before and a conserved {chr(955)}exp(e{chr(8320)}/17) after a "
        f"mid-twentieth-century change point. Using the same public data, we ask whether this is a "
        f"placeless law or partly an artefact of aligning asynchronous national transitions on a "
        f"common calendar, and whether it holds at sub-national scales. Country-specific onsets "
        f"span {onset['t_fert_min']}\u2013{onset['t_fert_max']} (IQR {onset['t_fert_iqr']:.0f} years). "
        f"Event-time alignment reduces between-country {chr(955)} variance by a factor of "
        f"{coll['variance_reduction_ratio']:.1f}, but tightens only the first conserved quantity. "
        f"Within countries, regional {chr(955)} at matched e{chr(8320)} differs as much as the "
        f"cross-national variation of the second pathway. The two pathways are a timing-locked "
        f"regularity plus heterogeneous convergence rather than a single universal constant."
    )
    doc.add_paragraph(abstract)

    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("demographic transition; fertility; life expectancy; universality; "
               "sub-national heterogeneity; event-time alignment")

    # ---- Body ----
    H(doc, "A cross-disciplinary claim", 2)
    p1 = (
        f"The search for law-like regularities in human populations has a long history "
        f"(Notestein, 1945; Kirk, 1996). Itao (2026) contributes a striking recent example: "
        f"plotting {chr(955)} against life expectancy at birth e{chr(8320)} for "
        f"{base['n_countries']} countries, the paper reports that national trajectories are "
        f"organised by two conserved quantities, {chr(955)}e{chr(8320)} before and "
        f"{chr(955)}exp(e{chr(8320)}/17) after a change point near the mid-twentieth century. "
        f"Importing statistical-physics search for conserved quantities into demography is a "
        f"valuable cross-disciplinary move, and the household model linking the shift to returns "
        f"to schooling (Galor, 2011) gives a plausible mechanism. We take the two-regime "
        f"description as a starting point and ask how far the regularity survives when timing "
        f"and spatial scale are treated more carefully."
    )
    doc.add_paragraph(p1)

    H(doc, "Asynchronous transitions, not a global turning point", 2)
    p2 = (
        f"If the reported change point were a genuine world-historical moment, most countries "
        f"should transition near it. They do not. Using Coale's (1973) onset definition\u2014the "
        f"first irreversible decline in a smoothed {chr(955)} series\u2014country-specific onsets "
        f"range from {onset['t_fert_min']} to {onset['t_fert_max']}, an inter-quartile range of "
        f"{onset['t_fert_iqr']:.0f} years; only {_pct(pco['frac_within_10y_of_1950'])} begin within "
        f"a decade of 1950 (Figure 1A). A pooled calendar split is sharply estimated at "
        f"{base['best_single_change_year']} (95% CI {int(boot['ci95_lo'])}\u2013"
        f"{int(boot['ci95_hi'])}), consistent with Itao (2026), but this narrowness is an "
        f"aggregation artefact: it summarises asynchronous national transitions rather than a "
        f"synchronised global event (Bongaarts & Watkins, 1996; Reher, 2004)."
    )
    doc.add_paragraph(p2)

    H(doc, "Event alignment sharpens the first pathway only", 2)
    p3 = (
        f"If transitions are asynchronous, a placeless law should look clearer when each "
        f"country is read on its own clock. Re-aligning trajectories on each country's onset "
        f"reduces the between-country standard deviation of {chr(955)} from "
        f"{coll['sd_lambda_calendar']:.1f} to {coll['sd_lambda_event']:.1f} births per 1,000 in "
        f"the {chr(177)}25-year window\u2014a {coll['variance_reduction_ratio']:.1f}-fold reduction "
        f"in variance (Figure 1B). Under alignment, the first conserved quantity, "
        f"{chr(955)}e{chr(8320)}, tightens from CV\u2009=\u2009{cvc['cv_lambda_e0_phaseI']:.2f} to "
        f"{cve['cv_lambda_e0_phaseI']:.2f}, a timing-locked regularity that calendar-time pooling "
        f"partly obscures. The second quantity, {chr(955)}exp(e{chr(8320)}/17), does not improve "
        f"(CV\u2009=\u2009{cvc['cv_lambda_exp_phaseII']:.2f} versus {cve['cv_lambda_exp_phaseII']:.2f}; Figure 1C). "
        f"The two pathways are asymmetric: the pre-transition process has a conservation-like "
        f"structure once timing is modelled, whereas the post-transition regime remains a loose, "
        f"heterogeneous convergence."
    )
    doc.add_paragraph(p3)

    caption = (
        f"Figure 1. (A) Fertility-transition onsets span "
        f"{onset['t_fert_max'] - onset['t_fert_min']} years; the shaded band is the 95% CI for the "
        f"pooled calendar change point. (B) Event-time alignment (each country at its own onset) "
        f"reduces the between-country SD of {chr(955)} by a factor of "
        f"{coll['variance_reduction_ratio']:.1f}. (C) Alignment lowers the CV of the first "
        f"conserved quantity ({chr(955)}e{chr(8320)}) but not the second "
        f"({chr(955)}exp(e{chr(8320)}/17)). (D) Within five decentralised states, regions at "
        f"nearly identical e{chr(8320)} scatter off the single 'universal' Phase II isocline, so the "
        f"country-level curve does not carry to sub-national scales."
    )
    add_figure(doc, os.path.join(FIGS_DIR, "fig_main.png"), caption, width=6.4)

    H(doc, "No single curve describes within-country variation", 2)
    p4 = (
        f"A genuinely placeless law should also hold below the national level. Region-level data "
        f"for {n_regions_total} regions of five decentralised states (Eurostat, 2024; National "
        f"Center for Health Statistics, 2022; U.S. Census Bureau, 2020) show that regions of the "
        f"same country at matched e{chr(8320)} differ substantially in {chr(955)} (Figure 1D). The "
        f"within-country CV of {chr(955)} among regions within a three-year e{chr(8320)} band ranges "
        f"from {matched_series.min():.2f} to {matched_series.max():.2f} (median "
        f"{matched_series.median():.2f}), comparable to the cross-national Phase II CV "
        f"({cve['cv_lambda_exp_phaseII']:.2f}). In Spain, regional {chr(955)} in {es['year']} "
        f"spans {es['lambda_min']:.1f}\u2013{es['lambda_max']:.1f} births per 1,000 despite an "
        f"e{chr(8320)} spread of only {es['e0_max'] - es['e0_min']:.1f} years. Longitudinal "
        f"sub-populations from the Human Mortality Database (2024) likewise diverge: "
        f"East and West Germany follow different "
        f"{chr(955)}\u2013e{chr(8320)} paths, with a mean matched-e{chr(8320)} gap of "
        f"{de['mean_abs_lambda_gap']:.1f} births per 1,000 ({_pct(de['mean_rel_gap'])} of the mean)."
    )
    doc.add_paragraph(p4)

    H(doc, "Conclusion", 2)
    p5 = (
        f"None of this overturns the regularity Itao (2026) documents. It does relocate it. The "
        f"first pathway is a timing-locked mortality\u2013fertility co-movement that becomes sharper "
        f"once each country is read on its own clock. The second pathway is better understood as "
        f"heterogeneous convergence toward low fertility than as a single universal constant. Event-time "
        f"alignment and hierarchical modelling are natural complements to the statistical-physics "
        f"approach, pushing the same reasoning down to the within-country scales where policy "
        f"operates."
    )
    doc.add_paragraph(p5)

    # ---- Required statements ----
    H(doc, "Acknowledgements", 2)
    doc.add_paragraph("The author thanks Kenji Itao for making the original analysis and data "
                      "openly available.")
    H(doc, "Financial Support", 2)
    doc.add_paragraph("This research received no specific grant.")
    H(doc, "Conflicts of Interest", 2)
    doc.add_paragraph("The author declares none.")

    # ---- References (APA, alphabetical) ----
    H(doc, "References", 2)
    refs = [
        "Bongaarts, J., & Watkins, S. C. (1996). Social interactions and contemporary fertility "
        "transitions. Population and Development Review, 22(4), 639\u2013682.",
        "Coale, A. J. (1973). The demographic transition reconsidered. In International "
        "Population Conference, Li\u00e8ge 1973 (Vol. 1, pp. 53\u201372). IUSSP.",
        "Eurostat. (2024). Regional demographic statistics [Data set]. "
        "https://ec.europa.eu/eurostat/web/regions/data/database",
        "Galor, O. (2011). Unified growth theory. Princeton University Press.",
        "Human Mortality Database. (2024). University of California, Berkeley, and Max Planck "
        "Institute for Demographic Research. https://www.mortality.org",
        "Itao, K. (2026). Two universal pathways in demographic transition. Evolutionary Human "
        "Sciences. https://doi.org/10.1017/ehs.2026.10054",
        "Kirk, D. (1996). Demographic transition theory. Population Studies, 50(3), 361\u2013387.",
        "National Center for Health Statistics. (2022). U.S. state life expectancy by sex, 2019 "
        "[Data set]. Centers for Disease Control and Prevention. https://data.cdc.gov",
        "Notestein, F. W. (1945). Population\u2014the long view. In T. W. Schultz (Ed.), Food for "
        "the world (pp. 36\u201357). University of Chicago Press.",
        "Reher, D. S. (2004). The demographic transition revisited as a global process. "
        "Population, Space and Place, 10(1), 19\u201341.",
        "U.S. Census Bureau. (2020). Population estimates program (PEP): Components of change "
        "[Data set]. https://www.census.gov/programs-surveys/popest.html",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)

    _add_line_numbers_and_page_numbers(doc)

    out = os.path.join(MANU_DIR, "EHS_commentary_demographic_transition.docx")
    doc.save(out)
    print("saved", out)

    total_text = abstract + " " + caption + " " + p1 + " " + p2 + " " + p3 + " " + p4 + " " + p5 + " " + " ".join(refs)
    print("word count (abstract+body+caption+refs):", _word_count(total_text))

if __name__ == "__main__":
    main()
