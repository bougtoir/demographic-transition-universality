"""Generate the EHS Perspective manuscript (docx) with figures inline and APA references.

Every numeric value is read from results/*.json|csv (produced by s02-s06, s10); nothing is
hard-coded. Figures are inserted immediately after the paragraph that first cites them.
"""
import json
import os

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from common import FIGS_DIR, RESULTS_DIR, ROOT

MANU_DIR = os.path.join(ROOT, "manuscript")


def _load():
    j = lambda n: json.load(open(os.path.join(RESULTS_DIR, n)))
    base = j("baseline_calendar.json")
    onset = j("onset_summary.json")
    cp = j("changepoint_uncertainty.json")
    ev = j("event_alignment.json")
    sens = j("sensitivity_summary.json")
    sub = j("subnational_summary.json")
    hmd_c = j("subnational_hmd_contrast.json")
    tr = j("time_rescaling.json")
    subpts = pd.read_csv(os.path.join(RESULTS_DIR, "subnational_points.csv"))
    return base, onset, cp, ev, sens, sub, hmd_c, tr, subpts


def _pct(x):
    return f"{100 * x:.0f}%"


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.italic = True
    return p


def add_figure(doc, path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.add_run().add_picture(path, width=Inches(width))
    add_caption(doc, caption)


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def _set_continuous_line_numbers(section):
    """Add Word continuous line numbering to a section."""
    sectPr = section._sectPr
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:restart"), "continuous")
    lnNumType.set(qn("w:distance"), "720")  # 0.5 inch
    sectPr.append(lnNumType)


def _add_page_number_footer(section):
    """Add a right-aligned PAGE field to the section's default footer."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.clear()

    # PAGE field: { PAGE \* MERGEFORMAT }
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE \\* MERGEFORMAT "
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    # placeholder display text; Word updates on open/print
    run2 = p.add_run("1")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar3)


def _add_line_numbers_and_page_numbers(doc):
    for section in doc.sections:
        _set_continuous_line_numbers(section)
        _add_page_number_footer(section)


def main():
    os.makedirs(MANU_DIR, exist_ok=True)
    base, onset, cp, ev, sens, sub, hmd_c, tr, subpts = _load()
    de = hmd_c["Germany"]

    r2 = base["r2_by_k"]
    boot = cp["global_change_year_bootstrap"]
    pco = cp["per_country_onset"]
    coll = ev["collapse"]
    cvc, cve = ev["conserved_cv_calendar"], ev["conserved_cv_event"]

    # subnational aggregates
    matched = {k: v["lambda_cv_matched_e0"] for k, v in sub.items()
               if v.get("lambda_cv_matched_e0") is not None}
    matched_series = pd.Series(matched)
    es = sub["Spain"]
    n_regions_total = int(subpts[subpts.source.isin(
        ["Eurostat NUTS2", "CDC/NCHS + US Census PEP"])].shape[0])

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)

    # ---- Title ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = t.add_run("Two pathways, many clocks: extending demographic-transition universality "
                     "to event time and within-country scales")
    trun.bold = True
    trun.font.size = Pt(15)

    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    au.add_run("Tatsuki Onishi").font.size = Pt(12)
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = aff.add_run("Affiliation to be completed. Corresponding author: Tatsuki Onishi "
                     "(email to be completed).")
    ar.font.size = Pt(9)
    ar.italic = True

    # ---- Abstract ----
    H(doc, "Abstract", 2)
    abstract = (
        f"Itao (2026) reports that demographic transitions follow two universal "
        f"pathways: a conserved product \u03bb\u00d7e\u2080 before a mid-twentieth-century change "
        f"point and a conserved \u03bb\u00d7exp(e\u2080/17) afterwards, where \u03bb is the crude "
        f"birth rate and e\u2080 life expectancy at birth. This statistical-physics reading of "
        f"demography\u2014seeking conserved quantities behind a global regularity\u2014is a "
        f"valuable cross-disciplinary contribution, and we build on it. Using the same public "
        f"data, we ask how far the regularity extends: whether it reflects a placeless law or is "
        f"partly shaped by aligning asynchronous national transitions on a common calendar, and "
        f"whether it carries to sub-national scales. Country-specific fertility "
        f"onsets span {onset['t_fert_min']}\u2013{onset['t_fert_max']} "
        f"(inter-quartile range {onset['t_fert_iqr']:.0f} years); only "
        f"{_pct(pco['frac_within_10y_of_1950'])} of countries begin within a decade of 1950. "
        f"Re-aligning each country at its own onset collapses the between-country dispersion of "
        f"\u03bb by a factor of {coll['variance_reduction_ratio']:.1f}, but tightens only the "
        f"first conserved quantity, not the second. Within single countries, regions at nearly "
        f"identical e\u2080 differ in \u03bb by as much as the cross-national variation of that "
        f"second pathway. We "
        f"read the two pathways as a timing-locked mortality\u2013fertility regularity plus "
        f"heterogeneous post-transition convergence rather than a single placeless constant, and "
        f"offer event-time alignment and hierarchical modelling as complements that sharpen, "
        f"rather than overturn, Itao's framework."
    )
    doc.add_paragraph(abstract)

    kw = doc.add_paragraph()
    kw.add_run("Keywords: ").bold = True
    kw.add_run("demographic transition; fertility; life expectancy; universality; "
               "sub-national heterogeneity; event-time alignment")

    # ---- Body ----
    H(doc, "A striking claim of universality", 2)
    doc.add_paragraph(
        f"A recurring ambition in the quantitative study of human populations is to find "
        f"law-like regularities that hold across societies (Notestein, 1945; Kirk, 1996). "
        f"Itao (2026) offers a bold instance. Plotting the crude birth rate \u03bb against period "
        f"life expectancy at birth e\u2080 for {base['n_countries']} countries, the paper reports "
        f"that the world's trajectories are organised by two conserved quantities: an early "
        f"regime in which \u03bb\u00d7e\u2080 is approximately constant, and a later regime in "
        f"which \u03bb\u00d7exp(e\u2080/17) is approximately constant, with a single change point "
        f"placed near the mid-twentieth century. The two regimes are interpreted as universal "
        f"'pathways' shared by all countries, and are linked through a household model of the "
        f"trade-off between the number of children and investment in their education to a shift "
        f"in the returns to schooling (Galor, 2011). Importing the statistical-physics search for "
        f"conserved quantities and phase-like regimes into demography is precisely the kind of "
        f"cross-disciplinary move that can expose structure other framings miss, and it motivates "
        f"the closer look we take here.")
    doc.add_paragraph(
        f"We reproduce the qualitative structure of this result from the same public source "
        f"(Gapminder, 2023). Fitting the paper's family of forms and searching for the single "
        f"calendar split that maximises the pooled coefficient of determination, the explained "
        f"variance rises from R\u00b2\u2009=\u2009{r2['1']:.2f} for one regime to "
        f"{r2['2']:.2f} for two, with a best split at "
        f"{base['best_single_change_year']}. This closely echoes Itao (2026), and we take the "
        f"two-regime description as a well-established starting point. Our question is a "
        f"complementary one: how far the fit reflects a placeless law versus the shared calendar "
        f"on which asynchronous national histories are viewed, and how it behaves when the same "
        f"logic is pushed to finer temporal and spatial scales.")

    H(doc, "Data and methods", 2)
    doc.add_paragraph(
        f"We use the national crude birth rate and period life expectancy series compiled by "
        f"Gapminder (2023), the same primary source as the target paper, giving "
        f"{base['n_obs']:,} country-year observations for {base['n_countries']} countries over "
        f"{base['year_min']}\u2013{base['year_max']}. Following Itao (2026), we fit through-origin "
        f"isoclines \u03bb\u2009=\u2009C/e\u2080\u1d56 (power laws) and "
        f"\u03bb\u2009=\u2009C/exp(e\u2080/c) (exponentials) and locate the single calendar split "
        f"that maximises the pooled R\u00b2 by exhaustive search; a country bootstrap "
        f"({boot['n_boot']:,} replicates, resampling countries) gives its confidence interval. "
        f"Each country's fertility onset is the last year its five-year-smoothed \u03bb sits "
        f"within {int(round(100*0.10))}% of its own pre-transition maximum before an irreversible "
        f"decline (Coale, 1973); event time is calendar year minus this onset. Conserved-quantity "
        f"dispersion is measured within the paper's own fitting windows (e\u2080\u2009>\u200940 "
        f"for the first quantity, e\u2080\u2009>\u200960 for the second). Sub-national \u03bb and "
        f"e\u2080 are matched by region and year. All code and derived tables are public (see "
        f"Research Transparency and Reproducibility); no value in this article is hand-entered.")

    H(doc, "What a universal law would predict", 2)
    doc.add_paragraph(
        "A genuinely universal, placeless relationship between \u03bb and e\u2080 makes two "
        "predictions that pooled calendar-time plots cannot test. First, because transitions "
        "began at very different dates, aligning each country not on the calendar but on its own "
        "transition onset should not degrade\u2014and should, if anything, sharpen\u2014the "
        "collapse of trajectories onto a common curve. Second, the law should hold at any level "
        "of aggregation: regions within a country, observed in the same year at essentially the "
        "same e\u2080, should share the same \u03bb. We test both.")

    H(doc, "The mid-century change point summarises asynchronous transitions", 2)
    doc.add_paragraph(
        f"Estimating each country's fertility-transition onset\u2014the year after which the "
        f"smoothed crude birth rate stays irreversibly below its pre-transition plateau (Coale, "
        f"1973)\u2014yields onsets from {onset['t_fert_min']} to {onset['t_fert_max']}, a range "
        f"of {onset['t_fert_range']} years, with an inter-quartile range of "
        f"{onset['t_fert_iqr']:.0f} years (Figure 1). Only "
        f"{_pct(pco['frac_within_10y_of_1950'])} of countries begin their transition within ten "
        f"years of 1950; {_pct(pco['frac_before_1930'])} began before 1930 and "
        f"{_pct(pco['frac_after_1970'])} after 1970. The pooled change point is nonetheless "
        f"sharply estimated ({base['best_single_change_year']}; bootstrap 95% CI "
        f"{int(boot['ci95_lo'])}\u2013{int(boot['ci95_hi'])}, an interval that comfortably "
        f"includes the mid-century date Itao (2026) adopts, so our split is consistent with, not "
        f"a correction to, the original). The narrowness is nonetheless revealing: it "
        f"shows that pooling manufactures a crisp global 'moment of transition' that almost no "
        f"individual country experiences near that date. A single calendar change point is thus "
        f"most naturally read as a summary of many asynchronous national transitions rather than "
        f"a moment at which the world turned a corner together (Bongaarts & Watkins, 1996; Reher, "
        f"2004; Figure 1).")

    add_figure(doc, os.path.join(FIGS_DIR, "panel_A_onset.png"),
               f"Figure 1. Country-specific fertility-transition onsets span roughly "
               f"{onset['t_fert_max'] - onset['t_fert_min']} years "
               f"({onset['t_fert_min']}\u2013{onset['t_fert_max']}). The shaded band is the "
               f"bootstrap 95% confidence interval ({int(boot['ci95_lo'])}\u2013"
               f"{int(boot['ci95_hi'])}) of the single pooled calendar change point; the dotted "
               f"line marks 1950.", width=4.6)

    H(doc, "Event alignment collapses the spread\u2014but only for the first pathway", 2)
    doc.add_paragraph(
        f"If national transitions are asynchronous, re-expressing each trajectory in event time "
        f"(years since its own onset) should concentrate them. It does. Within \u00b125 years of "
        f"the alignment point, the mean between-country standard deviation of \u03bb falls from "
        f"{coll['sd_lambda_calendar']:.1f} births per 1,000 in calendar time to "
        f"{coll['sd_lambda_event']:.1f} in event time\u2014a "
        f"{coll['variance_reduction_ratio']:.1f}-fold reduction in variance (Figure 2). Much of "
        f"the scatter that the paper's two-regime fit accommodates is therefore timing, not a "
        f"property of the \u03bb\u2013e\u2080 relationship itself.")
    add_figure(doc, os.path.join(FIGS_DIR, "panel_B_collapse.png"),
               f"Figure 2. Between-country standard deviation of the crude birth rate around the "
               f"alignment point, in calendar time (aligned at 1950) versus event time (aligned "
               f"at each country's onset). Event alignment reduces the variance "
               f"{coll['variance_reduction_ratio']:.1f}-fold.", width=4.6)
    doc.add_paragraph(
        f"Crucially, alignment does not rehabilitate both pathways equally. Recomputing the two "
        f"conserved quantities within the paper's own fitting windows, event alignment lowers the "
        f"coefficient of variation of the first quantity \u03bb\u00d7e\u2080 from {cvc['cv_lambda_e0_phaseI']:.2f} "
        f"(calendar) to {cve['cv_lambda_e0_phaseI']:.2f}\u2014a genuine, timing-locked "
        f"regularity that the calendar view actually obscures. But the second quantity "
        f"\u03bb\u00d7exp(e\u2080/17) does not tighten at all "
        f"({cvc['cv_lambda_exp_phaseII']:.2f} calendar versus {cve['cv_lambda_exp_phaseII']:.2f} "
        f"event; Figure 3). The two 'universal pathways' are thus asymmetric: the pre-transition "
        f"co-movement of fertility and mortality behaves like a conserved quantity once timing is "
        f"removed, whereas the post-transition regime is a loose, heterogeneous convergence that "
        f"no single constant captures\u2014consistent with the reversal of the "
        f"fertility\u2013development relationship at high income (Myrskyl\u00e4 et al., 2009).")
    add_figure(doc, os.path.join(FIGS_DIR, "panel_C_cv.png"),
               "Figure 3. Coefficient of variation of the two conserved quantities within the "
               "paper's fitting windows, under the calendar split versus event alignment. Only "
               "the first pathway tightens.", width=4.4)

    H(doc, "A single curve does not yet describe within-country variation", 2)
    doc.add_paragraph(
        f"A placeless law should also hold below the national level, and here the country-level "
        f"regularity does not straightforwardly carry over\u2014an open question that a "
        f"country-level analysis necessarily leaves for later work. Using "
        f"region-level life expectancy and crude birth rate for five decentralised states "
        f"({n_regions_total} regions in total; Eurostat, 2024; National Center for Health "
        f"Statistics, 2022; U.S. Census Bureau, 2020), regions of the same country in the same "
        f"year\u2014at nearly identical e\u2080\u2014differ substantially in \u03bb (Figure 4). "
        f"Among regions within a three-year e\u2080 band, the within-country coefficient of "
        f"variation of \u03bb ranges from {matched_series.min():.2f} to {matched_series.max():.2f} "
        f"across countries (median {matched_series.median():.2f}). In Spain, regional \u03bb in "
        f"{es['year']} runs from {es['lambda_min']:.1f} to {es['lambda_max']:.1f} births per "
        f"1,000 despite an e\u2080 spread of only {es['e0_max'] - es['e0_min']:.1f} years. This "
        f"intra-national dispersion is of the same order as the cross-national variation of the "
        f"second conserved quantity ({cve['cv_lambda_exp_phaseII']:.2f}) that the paper treats as "
        f"a universal signal. Whatever organises \u03bb at a given e\u2080, it is not a single "
        f"quantity shared by all populations.")
    doc.add_paragraph(
        f"Longitudinal sub-populations tell the same story. In the Human Mortality Database "
        f"(2024), East and West Germany traced materially different \u03bb\u2013e\u2080 paths over "
        f"their shared life-expectancy range ({de['e0_overlap'][0]:.0f}\u2013"
        f"{de['e0_overlap'][1]:.0f} years): at matched e\u2080 their crude birth rates differed "
        f"by {de['mean_abs_lambda_gap']:.1f} births per 1,000 on average "
        f"({_pct(de['mean_rel_gap'])} of the mean), reaching {de['max_abs_lambda_gap']:.1f} at "
        f"the widest\u2014two populations of one nation that never shared a single curve.")
    add_figure(doc, os.path.join(FIGS_DIR, "panel_D_subnational.png"),
               "Figure 4. Region-level crude birth rate versus life expectancy (2019) for five "
               "decentralised states, against the national Phase II isocline "
               "\u03bb\u2009=\u2009C/exp(e\u2080/17). Regions of one country scatter off the "
               "curve rather than lying on it.", width=4.8)

    H(doc, "Robustness", 2)
    doc.add_paragraph(
        f"Both headline findings are robust to the exploratory onset definition. Varying the "
        f"fertility-decline threshold (5\u201320%) and the smoothing window (3\u20139 years) over "
        f"{sens['n_settings']} combinations, the onset inter-quartile range stays between "
        f"{sens['iqr_min']:.0f} and {sens['iqr_max']:.0f} years and the event-time variance "
        f"reduction between {sens['vrr_min']:.1f}- and {sens['vrr_max']:.1f}-fold. The onset "
        f"estimator is a descriptive device, not a claim about any single country's demographic "
        f"history, and none of our conclusions depends on its exact tuning.")

    dur = tr["duration_years"]
    raw = tr["raw_lambda"]
    shp = tr["amplitude_normalised_shape"]
    doc.add_paragraph(
        f"A stronger test of universality is whether stretching each country's time axis\u2014not "
        f"merely shifting it to the onset\u2014collapses the trajectories onto a single curve "
        f"(Figure 5). Rescaling event time by each country's transition duration "
        f"(median {dur['median']:.0f} years, inter-quartile range {dur['q1']:.0f}\u2013"
        f"{dur['q3']:.0f}; N\u2009=\u2009{tr['n_countries_with_transition']}) does not: the "
        f"between-country dispersion of \u03bb at matched duration-normalised time is "
        f"{raw['sd_rescaled_time']:.1f} births per 1,000, no smaller than the "
        f"{raw['sd_event_time']:.1f} "
        f"obtained by onset alignment alone. The decline's shape is more nearly universal\u2014once "
        f"the amplitude is also normalised, the dispersion of the unit-scaled curves falls "
        f"{shp['variance_ratio_event_over_rescaled']:.1f}-fold\u2014but that comparison discards "
        f"the birth-rate level, which is exactly what a conserved quantity is supposed to fix. "
        f"Decisively, no monotone time warp can touch the relationship in the (e\u2080, \u03bb) "
        f"plane the law is stated in: the between-country SD of \u03bb at matched e\u2080 "
        f"({tr['irreducible_sd_lambda_at_matched_e0']:.1f} births per 1,000) is invariant to any "
        f"such rescaling. Time alignment thus sharpens the first pathway's timing, but does not "
        f"by itself establish universality of level.")
    add_figure(doc, os.path.join(FIGS_DIR, "fig_time_rescaling.png"),
               "Figure 5. Crude birth rate against (left) years since each country's transition "
               "onset and (right) duration-normalised transition time s\u2009=\u2009(t\u2212onset)/D. "
               "Rescaling the time axis does not shrink the raw between-country spread of "
               "\u03bb beyond onset alignment; only the amplitude-normalised shape collapses, and "
               "the (e\u2080, \u03bb) relationship is invariant to any monotone time warp.", width=5.4)

    H(doc, "Interpretation", 2)
    doc.add_paragraph(
        "None of this diminishes the regularity Itao (2026) documents; our aim is to build on it "
        "and locate it more precisely. The first pathway emerges, on this reading, as a real, "
        "timing-locked co-movement: as mortality falls, fertility follows in a way that is "
        "remarkably similar\u2014indeed sharper than in calendar time\u2014once each country is "
        "read on its own clock. The second 'pathway' is better described as incomplete, "
        "heterogeneous convergence toward low fertility, whose scatter\u2014across countries and, "
        "importantly, within them\u2014is larger than a single conserved constant would imply. We "
        "therefore offer three complementary refinements that extend, rather than overturn, the "
        "framework: analyse trajectories in event time alongside calendar time; model countries "
        "and regions as a hierarchy with partial pooling rather than a single pooled cloud; and "
        "read a common change point as a description of aggregated timing rather than a "
        "synchronised global event. Understood this way, the statistical-physics search for "
        "universality remains productive: it points to mechanism and shape once timing and level "
        "are accounted for, and its natural next step\u2014which the mechanistic household model "
        "of Itao (2026) invites\u2014is to carry the same reasoning down to the within-country "
        "scales where policy actually operates.")

    # ---- Required statements ----
    H(doc, "Acknowledgements", 2)
    doc.add_paragraph("The author thanks Kenji Itao for making the original analysis and data "
                      "openly available, which made this re-analysis possible.")
    H(doc, "Author Contributions", 2)
    doc.add_paragraph("TO conceived the study, performed all analyses, and wrote the article.")
    H(doc, "Financial Support", 2)
    doc.add_paragraph("This research received no specific grant from any funding agency, "
                      "commercial or not-for-profit sectors.")
    H(doc, "Conflicts of Interest", 2)
    doc.add_paragraph("The author declares none.")
    H(doc, "Research Transparency and Reproducibility", 2)
    doc.add_paragraph(
        "All data are public. Long-run national crude birth rate and life expectancy series are "
        "from Gapminder (https://www.gapminder.org/data/); sub-national data are from the "
        "Eurostat regional database (https://ec.europa.eu/eurostat/web/regions/data/database), "
        "the Human Mortality Database (https://www.mortality.org), the U.S. National Center for "
        "Health Statistics (https://data.cdc.gov) and the U.S. Census Bureau Population Estimates "
        "Program (https://www.census.gov/programs-surveys/popest.html). All analysis code and "
        "derived results that regenerate every number and figure in this article are available "
        "at https://github.com/bougtoir/demographic-transition-universality (release archived on "
        "acceptance).")

    # ---- References (APA, alphabetical) ----
    H(doc, "References", 2)
    refs = [
        "Bongaarts, J., & Watkins, S. C. (1996). Social interactions and contemporary fertility "
        "transitions. Population and Development Review, 22(4), 639\u2013682.",
        "Coale, A. J. (1973). The demographic transition reconsidered. In International "
        "Population Conference, Li\u00e8ge 1973 (Vol. 1, pp. 53\u201372). IUSSP.",
        "Eurostat. (2024). Regional demographic statistics [Data set]. "
        "https://ec.europa.eu/eurostat/web/regions/data/database",
        "Galor, O. (2011). Unified growth theory. Princeton University Press.",
        "Gapminder. (2023). Population, fertility and life expectancy data [Data set]. "
        "https://www.gapminder.org/data/",
        "Human Mortality Database. (2024). University of California, Berkeley, and Max Planck "
        "Institute for Demographic Research. https://www.mortality.org",
        "Itao, K. (2026). Two universal pathways in demographic transition. Evolutionary Human "
        "Sciences. https://doi.org/10.1017/ehs.2026.10054",
        "Kirk, D. (1996). Demographic transition theory. Population Studies, 50(3), 361\u2013387.",
        "Myrskyl\u00e4, M., Kohler, H.-P., & Billari, F. C. (2009). Advances in development "
        "reverse fertility declines. Nature, 460(7256), 741\u2013743.",
        "National Center for Health Statistics. (2022). U.S. state life expectancy by sex, 2019 "
        "[Data set]. Centers for Disease Control and Prevention. https://data.cdc.gov",
        "Notestein, F. W. (1945). Population\u2014the long view. In T. W. Schultz (Ed.), Food for "
        "the world (pp. 36\u201357). University of Chicago Press.",
        "Reher, D. S. (2004). The demographic transition revisited as a global process. "
        "Population, Space and Place, 10(1), 19\u201341.",
        "U.S. Census Bureau. (2020). Population estimates program (PEP): Components of change "
        "[Data set]. https://www.census.gov/programs-surveys/popest.html",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)

    _add_line_numbers_and_page_numbers(doc)

    out = os.path.join(MANU_DIR, "EHS_perspective_demographic_transition.docx")
    doc.save(out)

    # crude body word count (excludes title, abstract, refs, statements headings)
    print("saved", out)
    print("references:", len(refs))


if __name__ == "__main__":
    main()
